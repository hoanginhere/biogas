import os
from googletrans import Translator
from docx import Document
from tqdm import tqdm

def translate_text(text, src='en', dest='vi'):
    translator = Translator()
    translated = translator.translate(text, src=src, dest=dest)
    return translated.text

def translate_docx(file_path, output_path, chunk_size=100):
    doc = Document(file_path)
    new_doc = Document()

    for paragraph in tqdm(doc.paragraphs, desc="Translating paragraphs"):
        text = paragraph.text
        if not text.strip():
            new_doc.add_paragraph('')
            continue
        
        words = text.split()
        translated_paragraphs = []
        for i in range(0, len(words), chunk_size):
            chunk = ' '.join(words[i:i+chunk_size])
            translated_chunk = translate_text(chunk)
            translated_paragraphs.append(translated_chunk)
        
        translated_text = ' '.join(translated_paragraphs)
        new_doc.add_paragraph(translated_text)
    
    new_doc.save(output_path)

if __name__ == "__main__":
    input_file = r"C:\Users\admin\Downloads\CAT I _ 2016 (1).docx"
    output_file = r"C:\Users\admin\Downloads\CAT I _ 2016 (1)_translated.docx"
    
    translate_docx(input_file, output_file)
